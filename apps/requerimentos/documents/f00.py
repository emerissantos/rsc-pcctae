from __future__ import annotations

from collections import defaultdict
from collections.abc import Iterable
from decimal import Decimal
from io import BytesIO

from django.db.models import Prefetch
from django.utils import timezone
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from apps.requerimentos.models import DocumentoLancamento, LancamentoItem, Requerimento

GREEN = "1F6339"
LIGHT_GREEN = "EAF3EC"
BORDER_GREEN = "9CB9A5"
TEXT_GRAY = RGBColor(74, 74, 74)
PLACEHOLDER = "[PREENCHER]"

CRITERIOS_OFICIAIS = {
    "I": (
        "Participação em grupos de trabalho, comissões, comitês, núcleos, "
        "representações ou similares"
    ),
    "II": "Projetos institucionais, gestão, ensino, pesquisa, extensão, inovação ou assistência",
    "III": "Premiações e reconhecimentos públicos",
    "IV": "Responsabilidades técnico-administrativas e/ou especializadas",
    "V": "Funções ou cargos de direção e assessoramento institucional",
    "VI": "Produção, prospecção e difusão de conhecimento",
}


def _fmt_decimal(value: Decimal | int | str | None) -> str:
    if value is None:
        return ""
    number = Decimal(str(value)).quantize(Decimal("0.01"))
    return f"{number:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _fmt_date(value) -> str:
    return value.strftime("%d/%m/%Y") if value else PLACEHOLDER


def _safe_text(value, fallback: str = PLACEHOLDER) -> str:
    text = str(value or "").strip()
    return text or fallback


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_borders(cell, *, color: str = BORDER_GREEN, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table, widths) -> None:
    table.autofit = False
    total_twips = sum(int(width.twips) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid_columns = list(table._tbl.tblGrid.gridCol_lst)
    for grid_column, width in zip(grid_columns, widths, strict=True):
        grid_column.set(qn("w:w"), str(int(width.twips)))

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_end])


def _apply_document_defaults(document: Document, requerimento: Requerimento) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor.from_string(GREEN)

    props = document.core_properties
    props.title = f"F-00 - Requerimento {requerimento.numero}"
    props.subject = "Formulário-padrão de requerimento do RSC-PCCTAE"
    props.author = "UFSB / CRSC-PCCTAE"
    props.comments = "Documento editável gerado pelo Sistema RSC-PCCTAE."

    for current_section in document.sections:
        footer = current_section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("UFSB / CRSC-PCCTAE - F-00 editável - Página ")
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = TEXT_GRAY
        _add_page_number(footer)


def _title(document: Document, text: str, *, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {min(level, 2)}"]
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(4)
    _set_keep_with_next(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12 if level == 1 else 10.5)
    run.font.color.rgb = RGBColor.from_string(GREEN)


def _add_intro(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("F-00 - FORMULÁRIO-PADRÃO DE REQUERIMENTO DO RSC-PCCTAE")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(GREEN)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(
        "Uso: servidor requerente - transcrição editável do formulário-padrão. "
        "Revise e complemente as informações antes de assinar ou protocolar."
    )
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = TEXT_GRAY

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(
        "Modelo para instrução de processo de Reconhecimento de Saberes e Competências "
        "(RSC-PCCTAE)"
    )
    run.bold = True
    run.font.size = Pt(10)


def _add_key_value_table(document: Document, rows: Iterable[tuple[str, str]]) -> None:
    rows = list(rows)
    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(5.4), Cm(12.0)]
    _set_table_widths(table, widths)
    for row, (label, value) in zip(table.rows, rows, strict=True):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_borders(cell)
            _set_cell_margins(cell)
        _set_cell_shading(row.cells[0], LIGHT_GREEN)
        label_p = row.cells[0].paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_run = label_p.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)
        value_p = row.cells[1].paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        value_run.font.size = Pt(9)
        if value.startswith("["):
            value_run.font.color.rgb = RGBColor(160, 76, 0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _marked_options(options: list[tuple[str, bool]]) -> str:
    return "   ".join(f"({'X' if selected else ' '}) {label}" for label, selected in options)


def _document_labels(lancamentos: list[LancamentoItem]) -> dict[int, list[str]]:
    labels: dict[int, list[str]] = defaultdict(list)
    sequence = 1
    for lancamento in lancamentos:
        for documento in lancamento.documentos.all():
            if not documento.ativo:
                continue
            labels[lancamento.pk].append(f"DOC-{sequence:03d} - {documento.nome_original}")
            sequence += 1
    return labels


def _criterion_table(
    document: Document,
    *,
    codigo: str,
    titulo: str,
    lancamentos: list[LancamentoItem],
    document_labels: dict[int, list[str]],
) -> Decimal:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(3)
    _set_keep_with_next(heading)
    run = heading.add_run(f"Critério {codigo} - {titulo}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)

    blank_rows = 2 if not lancamentos else 1
    table = document.add_table(rows=1 + len(lancamentos) + blank_rows + 1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(1.35), Cm(5.55), Cm(2.65), Cm(2.2), Cm(2.8), Cm(3.8)]
    _set_table_widths(table, widths)
    headers = [
        "Nº do item",
        "Critério específico",
        "Unidade de medida",
        "Pontuação",
        "Pontuação obtida",
        "Documentos comprobatórios (anexos)",
    ]
    header = table.rows[0]
    _set_repeat_table_header(header)
    header_values = zip(header.cells, headers, widths, strict=True)
    for index, (cell, header_text, width) in enumerate(header_values):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, GREEN)
        _set_cell_borders(cell, color="FFFFFF", size="4")
        _set_cell_margins(cell, top=65, start=65, bottom=65, end=65)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header_text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(7.2 if index in {1, 5} else 7.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    subtotal = Decimal("0.00")
    row_index = 1
    for lancamento in lancamentos:
        subtotal += lancamento.pontuacao_declarada
        docs = "\n".join(document_labels.get(lancamento.pk, [])) or "[SEM DOCUMENTO LISTADO]"
        quantity_unit = (
            f"{_fmt_decimal(lancamento.quantidade_declarada)} x "
            f"{lancamento.item_unidade_snapshot}"
        )
        description = lancamento.item_descricao_snapshot
        if lancamento.observacao:
            description += f"\nObservação do requerente: {lancamento.observacao}"
        values = [
            lancamento.item_codigo_snapshot,
            description,
            quantity_unit,
            _fmt_decimal(lancamento.item_pontos_snapshot),
            _fmt_decimal(lancamento.pontuacao_declarada),
            docs,
        ]
        row = table.rows[row_index]
        for cell, value, width in zip(row.cells, values, widths, strict=True):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_borders(cell)
            _set_cell_margins(cell, top=55, start=65, bottom=55, end=65)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            r.font.name = "Arial"
            r.font.size = Pt(7.6)
        row_index += 1

    for _ in range(blank_rows):
        row = table.rows[row_index]
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_borders(cell)
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(" ").font.size = Pt(8)
        row_index += 1

    subtotal_row = table.rows[row_index]
    merged = subtotal_row.cells[0].merge(subtotal_row.cells[3])
    merged.text = "Subtotal"
    _set_cell_shading(merged, LIGHT_GREEN)
    for cell in subtotal_row.cells:
        _set_cell_borders(cell)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    subtotal_row.cells[4].text = _fmt_decimal(subtotal)
    subtotal_row.cells[5].text = ""
    for index, cell in enumerate((merged, subtotal_row.cells[4])):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT if index == 1 else WD_ALIGN_PARAGRAPH.LEFT
        )
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
    return subtotal


def _add_declaration(document: Document, total: Decimal, nivel: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    text = (
        f"À vista das informações apresentadas, totalizo {_fmt_decimal(total)} pontos e atendo "
        f"aos critérios legais e regulamentares para o nível {nivel} do RSC-PCCTAE. "
        "Solicito a análise pela CRSC-PCCTAE."
    )
    p.add_run(text).font.size = Pt(10)

    _title(document, "4. Declaração de Conformidade Legal")
    p = document.add_paragraph(
        "Declaro, para os fins previstos no Decreto regulamentador do RSC-PCCTAE, que:"
    )
    p.paragraph_format.space_after = Pt(6)
    declarations = [
        "I - Todos os fatos apresentados ocorreram no exercício do cargo;",
        "II - Nenhuma atividade aqui declarada foi utilizada em requerimentos anteriores;",
        (
            "III - Toda a documentação anexada é autêntica e comprova integralmente "
            "as atividades apresentadas; e"
        ),
        (
            "IV - Tenho ciência de que informações falsas implicam responsabilidade "
            "administrativa, civil e penal."
        ),
    ]
    for text in declarations:
        p = document.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(text).font.size = Pt(10)

    document.add_paragraph().paragraph_format.space_after = Pt(15)
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(table, [Cm(10.5), Cm(6.5)])
    table.cell(0, 0).text = "____________________________________________"
    table.cell(0, 1).text = "____/____/________"
    table.cell(1, 0).text = "Assinatura do(a) servidor(a)"
    table.cell(1, 1).text = "Data"
    for row in table.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8.5)


def gerar_f00_docx(requerimento: Requerimento) -> bytes:
    """Gera o F-00 em DOCX editável a partir do estado atual do requerimento."""

    documentos_ativos = DocumentoLancamento.objects.filter(ativo=True).order_by(
        "created_at", "pk"
    )
    lancamentos = list(
        requerimento.lancamentos.select_related("item__requisito")
        .prefetch_related(Prefetch("documentos", queryset=documentos_ativos))
        .order_by("item__requisito__ordem", "item__ordem", "item_codigo_snapshot")
    )
    by_requirement: dict[str, list[LancamentoItem]] = defaultdict(list)
    for lancamento in lancamentos:
        by_requirement[lancamento.item.requisito.codigo].append(lancamento)

    document = Document()
    _apply_document_defaults(document, requerimento)
    _add_intro(document)

    vinculo = requerimento.vinculo
    servidor = vinculo.servidor
    pessoa = servidor.pessoa
    email = (
        servidor.email_atual
        or pessoa.email_institucional
        or requerimento.requerente.email
    )
    siape = vinculo.siape
    if vinculo.digito_siape:
        siape = f"{siape}-{vinculo.digito_siape}"

    _title(document, "1. Identificação do Servidor")
    _add_key_value_table(
        document,
        [
            (
                "Nome:",
                _safe_text(servidor.nome_atual or pessoa.nome or requerimento.requerente),
            ),
            ("Siape:", _safe_text(siape)),
            ("Cargo:", _safe_text(vinculo.cargo_nome)),
            (
                "Data de ingresso em Instituição Federal de Ensino:",
                _fmt_date(vinculo.data_admissao),
            ),
            (
                "Lotação:",
                _safe_text(vinculo.lotacao_nome or vinculo.unidade_exercicio_nome),
            ),
            ("Função/Encargo (se houver):", "[PREENCHER, SE HOUVER]"),
            ("Telefone/E-mail:", f"Telefone: {PLACEHOLDER} | E-mail: {_safe_text(email)}"),
        ],
    )

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Nível de Classificação:  A (   )   B (   )   C (   )   D (   )   E (   )")
    r.bold = True
    r.font.size = Pt(9.5)

    _title(document, "2. Informações do Requerimento")
    nivel_codigo = requerimento.nivel_pretendido.codigo.upper().strip()
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "Nível de RSC pretendido:  "
        + _marked_options(
            [(f"RSC-{code}", code == nivel_codigo) for code in CRITERIOS_OFICIAIS]
        )
    )
    r.bold = True
    r.font.size = Pt(9.3)

    total = requerimento.pontuacao_declarada or Decimal("0.00")
    minimum = requerimento.nivel_pretendido.pontuacao_minima or Decimal("0.00")
    excess = max(total - minimum, Decimal("0.00"))
    _add_key_value_table(
        document,
        [
            ("Pontuação mínima necessária:", _fmt_decimal(minimum)),
            ("Pontuação total apresentada:", _fmt_decimal(total)),
            ("Quantidade de critérios específicos utilizados:", str(len(lancamentos))),
            ("Pontuação total excedente (banco de pontos):", _fmt_decimal(excess)),
            ("Saldo de pontuação de concessão anterior:", "[PREENCHER, SE HOUVER]"),
            (
                "Número do processo relativo à concessão anterior do RSC-PCCTAE (se houver):",
                "[PREENCHER, SE HOUVER]",
            ),
        ],
    )

    document.add_page_break()
    _title(document, "3. Descrição das Atividades por Requisito Legal")
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "Os itens abaixo foram preenchidos a partir do requerimento eletrônico. "
        "Revise a descrição, a quantidade, a pontuação e a relação dos comprovantes, "
        "complementando o que for necessário."
    )
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = TEXT_GRAY

    document_labels = _document_labels(lancamentos)
    subtotals: dict[str, Decimal] = {}
    for code in ("I", "II", "III"):
        titulo = CRITERIOS_OFICIAIS[code]
        subtotals[code] = _criterion_table(
            document,
            codigo=code,
            titulo=titulo,
            lancamentos=by_requirement.get(code, []),
            document_labels=document_labels,
        )

    document.add_page_break()
    for code in ("IV", "V", "VI"):
        titulo = CRITERIOS_OFICIAIS[code]
        subtotals[code] = _criterion_table(
            document,
            codigo=code,
            titulo=titulo,
            lancamentos=by_requirement.get(code, []),
            document_labels=document_labels,
        )

    separator = document.add_paragraph()
    separator.paragraph_format.space_after = Pt(1)
    total_table = document.add_table(rows=1, cols=2)
    total_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(total_table, [Cm(14.8), Cm(3.6)])
    total_table.cell(0, 0).text = (
        "(Critério I + Critério II + Critério III + Critério IV + Critério V + "
        "Critério VI) TOTAL"
    )
    total_table.cell(0, 1).text = _fmt_decimal(sum(subtotals.values(), Decimal("0.00")))
    for cell in total_table.rows[0].cells:
        _set_cell_borders(cell)
        _set_cell_shading(cell, LIGHT_GREEN)
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(8.5)

    document.add_page_break()
    _add_declaration(document, total, requerimento.nivel_pretendido.nome)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Documento gerado em {timezone.localtime().strftime('%d/%m/%Y às %H:%M')} a partir do "
        f"requerimento {requerimento.numero}. Campos entre colchetes devem ser "
        "conferidos e preenchidos."
    )
    r.italic = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = TEXT_GRAY

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
