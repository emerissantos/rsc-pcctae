from datetime import date
from decimal import Decimal
from io import BytesIO

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from docx import Document

from apps.auditoria.models import EventoAuditoria
from apps.contas.models import IdentidadeExterna, Usuario
from apps.pessoas.models import PessoaInstitucional, Servidor, VinculoFuncional
from apps.pontuacao.models import ItemPontuacao, NivelRSC, Requisito
from apps.requerimentos.documents import gerar_f00_docx
from apps.requerimentos.models import DocumentoLancamento, LancamentoItem, Requerimento


@pytest.fixture
def requerimento_f00(db, tmp_path, settings):
    settings.MEDIA_ROOT = tmp_path
    usuario = Usuario.objects.create_user(
        username="maria",
        password="teste",
        email="maria@ufsb.edu.br",
        nome_exibicao="Maria da Silva",
    )
    pessoa = PessoaInstitucional.objects.create(
        id_institucional=900,
        nome="Maria da Silva",
        email_institucional="maria@ufsb.edu.br",
    )
    IdentidadeExterna.objects.create(
        usuario=usuario,
        pessoa=pessoa,
        id_usuario_externo=901,
        id_institucional=900,
        login="maria",
    )
    servidor = Servidor.objects.create(
        pessoa=pessoa,
        nome_atual="Maria da Silva",
        email_atual="maria@ufsb.edu.br",
    )
    vinculo = VinculoFuncional.objects.create(
        servidor=servidor,
        id_servidor_externo=902,
        siape="1234567",
        digito_siape="8",
        cargo_nome="Assistente em Administração",
        lotacao_nome="Pró-Reitoria de Gestão para Pessoas",
        unidade_exercicio_nome="PROGEPE",
        data_admissao=date(2020, 2, 3),
    )
    nivel = NivelRSC.objects.create(
        codigo="IV",
        nome="RSC-PCCTAE IV",
        pontuacao_minima=Decimal("30.00"),
        quantidade_minima_itens=1,
    )
    requisito_i = Requisito.objects.create(
        codigo="I",
        nome="Participação em comissões",
        ordem=1,
    )
    requisito_vi = Requisito.objects.create(
        codigo="VI",
        nome="Produção de conhecimento",
        ordem=6,
    )
    item_i = ItemPontuacao.objects.create(
        requisito=requisito_i,
        codigo="I.2",
        descricao="Presidência de comissão formalmente instituída",
        unidade="Por designação",
        pontos_por_quantidade=Decimal("4.50"),
    )
    item_vi = ItemPontuacao.objects.create(
        requisito=requisito_vi,
        codigo="VI.1",
        descricao="Produção de material técnico institucional",
        unidade="Por produto",
        pontos_por_quantidade=Decimal("7.50"),
    )
    requerimento = Requerimento.objects.create(
        requerente=usuario,
        vinculo=vinculo,
        nivel_pretendido=nivel,
        observacao_geral="Requerimento de teste.",
    )
    lancamento_i = LancamentoItem.objects.create(
        requerimento=requerimento,
        item=item_i,
        quantidade_declarada=Decimal("2"),
        observacao="Comissão de 2024 e comissão de 2025.",
    )
    lancamento_vi = LancamentoItem.objects.create(
        requerimento=requerimento,
        item=item_vi,
        quantidade_declarada=Decimal("3"),
    )
    DocumentoLancamento.objects.create(
        lancamento=lancamento_i,
        arquivo=SimpleUploadedFile("portaria-comissao.pdf", b"arquivo"),
        nome_original="portaria-comissao.pdf",
        tipo_mime="application/pdf",
    )
    DocumentoLancamento.objects.create(
        lancamento=lancamento_vi,
        arquivo=SimpleUploadedFile("manual-tecnico.docx", b"arquivo"),
        nome_original="manual-tecnico.docx",
        tipo_mime=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )
    requerimento.refresh_from_db()
    return usuario, requerimento


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.mark.django_db
def test_gerador_f00_preenche_dados_itens_e_documentos(requerimento_f00):
    _, requerimento = requerimento_f00

    content = gerar_f00_docx(requerimento)
    document = Document(BytesIO(content))
    text = _document_text(document)

    assert content.startswith(b"PK")
    assert "F-00 - FORMULÁRIO-PADRÃO DE REQUERIMENTO DO RSC-PCCTAE" in text
    assert "Maria da Silva" in text
    assert "1234567-8" in text
    assert "Assistente em Administração" in text
    assert "RSC-IV" in text
    assert "I.2" in text
    assert "VI.1" in text
    assert "portaria-comissao.pdf" in text
    assert "manual-tecnico.docx" in text
    assert "31,50" in text
    assert "[PREENCHER, SE HOUVER]" in text
    assert all(f"Critério {code} -" in text for code in ("I", "II", "III", "IV", "V", "VI"))


@pytest.mark.django_db
def test_requerente_baixa_f00_e_evento_e_auditado(client, requerimento_f00):
    usuario, requerimento = requerimento_f00
    client.force_login(usuario)

    response = client.get(reverse("requerimentos:gerar-f00", args=[requerimento.uuid]))

    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    assert response["Content-Type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "F-00-RSC-PCCTAE" in response["Content-Disposition"]
    assert response["Cache-Control"] == "private, no-store, no-cache, max-age=0"
    event = EventoAuditoria.objects.get(tipo=EventoAuditoria.Tipo.FORMULARIO_F00_GERADO)
    assert event.ator == usuario
    assert event.usuario_afetado == usuario
    assert event.objeto_id == str(requerimento.uuid)
    assert event.dados["quantidade_itens"] == 2


@pytest.mark.django_db
def test_outro_usuario_nao_gera_f00(client, requerimento_f00):
    _, requerimento = requerimento_f00
    outro = Usuario.objects.create_user(username="outro-f00", password="teste")
    client.force_login(outro)

    response = client.get(reverse("requerimentos:gerar-f00", args=[requerimento.uuid]))

    assert response.status_code == 403
    assert not EventoAuditoria.objects.filter(
        tipo=EventoAuditoria.Tipo.FORMULARIO_F00_GERADO
    ).exists()
